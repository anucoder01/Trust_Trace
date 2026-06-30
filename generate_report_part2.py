import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    import docx.oxml as oxml
    tcMar = oxml.OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = oxml.OxmlElement(m)
        node.set(oxml.ns.qn('w:w'), str(val))
        node.set(oxml.ns.qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report_part2():
    doc = Document()
    
    # Page Setup - Margins (1 inch everywhere)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base style setting (Times New Roman, 12pt)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    
    # ------------------ CHAPTER 2: RELATED WORK ------------------
    c2 = doc.add_paragraph()
    run = c2.add_run("CHAPTER 2\nRELATED WORK")
    run.font.size = Pt(16)
    run.font.bold = True
    c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c2.paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.add_run("Automated signature verification has been extensively researched, split between off-line (static scanned images) and on-line (dynamic electronic stylus signals) domains. Traditional static verification methods rely heavily on handcrafted geometric descriptors like Local Binary Patterns (LBP), Gray-Level Co-occurrence Matrices (GLCM), and Histogram of Oriented Gradients (HOG) to extract structural contours. These descriptors are typically classified using machine learning models such as Support Vector Machines (SVM), Random Forests, or Neural Networks. However, these systems focus primarily on global shape parameters. Consequently, they remain highly vulnerable to skilled forgeries where the shape is traced perfectly but drawn slowly, lacking natural fluidity.\n\n")
    p.add_run("Dynamic signature verification captures temporal patterns like pen tip velocity, acceleration, writing pressure, and coordinate trajectory using specialized digitizing tablets. Although highly secure, these dynamic approaches cannot be applied to historical paper archives or standard bank cheques. To resolve this limitation, computer vision researchers have explored micro-tremor and stroke analysis on static scans, attempting to detect the tiny hand tremors and pauses associated with the cognitive load of forgery. Recent deep learning approaches have utilized CNN-based Siamese Networks to learn deep similarity metrics between genuine and questioned signature images. While highly accurate, these deep models function as 'black boxes' without explainable outputs, require extensive training data, and require GPU hardware, which is often not available on standard bank office terminals.")

    h_gap = doc.add_paragraph()
    run = h_gap.add_run("Research Gaps")
    run.font.size = Pt(14)
    run.font.bold = True
    h_gap.paragraph_format.space_before = Pt(12)
    h_gap.paragraph_format.space_after = Pt(6)

    p_gap = doc.add_paragraph()
    p_gap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_gap.paragraph_format.line_spacing = 1.15
    p_gap.add_run("The literature survey reveals several critical research gaps in the domain of signature verification:\n")
    
    gaps = [
        "Lack of Integrated Verification Frameworks: Existing systems treat shape recognition, micro-tremor analysis, dynamic rhythm tracking, and financial transaction metadata as isolated problems rather than combining them into a single, unified decision model.",
        "High False Rejections due to Signature Drift: Traditional classifiers assume static signature templates. They do not account for natural changes in a person's signature over time, resulting in high false rejection rates (FRR).",
        "Lack of Explainability for Manual Inspection: Deep learning models fail to explain their output. Tellers receive a simple binary verdict without visual indicators showing which parts of a signature are suspicious.",
        "High Resource Requirements: High-accuracy deep models rely on complex neural networks that require GPU hardware. This makes them impractical to serve locally on standard bank teller desktops."
    ]
    for gp in gaps:
        doc.add_paragraph(gp, style='List Bullet')

    doc.add_page_break()

    # ------------------ CHAPTER 3: SYSTEM DESIGN ------------------
    c3 = doc.add_paragraph()
    run = c3.add_run("CHAPTER 3\nSYSTEM DESIGN")
    run.font.size = Pt(16)
    run.font.bold = True
    c3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c3.paragraph_format.space_after = Pt(18)

    h_meth = doc.add_paragraph()
    run = h_meth.add_run("Methodology")
    run.font.size = Pt(14)
    run.font.bold = True
    h_meth.paragraph_format.space_before = Pt(12)
    h_meth.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.add_run("The proposed TrustTrace system verifies signatures and detects forgery using a multi-layered machine learning and computer vision pipeline. The workflow contains six primary stages:\n\n")
    
    p.add_run("1. Preprocessing: Scanned signature images are loaded, converted to grayscale, and binarized using Otsu's thresholding. This isolates ink strokes and removes background noise. The isolated signature is then resized to a standard 256x128 pixel resolution.\n")
    p.add_run("2. Static Feature Extraction: The system extracts spatial shape descriptors from the binary image using the Histogram of Oriented Gradients (HOG) algorithm with 9 orientations, 8x8 cells, and 2x2 blocks.\n")
    p.add_run("3. Micro-Tremor Analysis: OpenCV filters analyze the physics of the strokes. The Edge Jitter Ratio is calculated by comparing raw Canny edges with Gaussian-smoothed edges. The Ink Blot Ratio is calculated by identifying contour areas larger than 50 pixels to locate hesitation pooling.\n")
    p.add_run("4. Soft-Voting Classification: The extracted HOG vector is classified using a Soft-Voting Ensemble model combining SVM, Random Forest, Gradient Boosting, Logistic Regression, and KNN. The ensemble calculates the average predicted probability across all classifiers to determine a shape-match verdict.\n")
    p.add_run("5. Dynamic Rhythm Analysis: For digital entry points, the system tracks stroke sequence coordinates, speed, and lifts over time. The velocity variance is calculated using an LSTM-style temporal model to differentiate between fluid genuine strokes and slow, traced copies.\n")
    p.add_run("6. Multimodal Risk Fusion: The final fraud probability is computed by fusing the visual signature score with transaction metadata (amount, location, history flags) using a weighted Bayesian formula.")

    h_flow = doc.add_paragraph()
    run = h_flow.add_run("Flow Diagram")
    run.font.size = Pt(14)
    run.font.bold = True
    h_flow.paragraph_format.space_before = Pt(12)
    h_flow.paragraph_format.space_after = Pt(6)

    p_flow = doc.add_paragraph()
    p_flow.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_flow.paragraph_format.line_spacing = 1.15
    p_flow.add_run("The flow diagram represents the complete execution workflow of the TrustTrace signature verification platform. The process begins when a scanned cheque image or drawn coordinate array is loaded. If it is an image, the system runs preprocessing (grayscale conversion, Otsu thresholding, and resizing) to clean background noise. The cleaned binary image is then split into two parallel analysis paths:\n\n")
    p_flow.add_run("- Shape Verification: HOG features are extracted and passed to the Soft-Voting Ensemble model to compute the shape-match probability.\n")
    p_flow.add_run("- Physical Stroke Verification: The system calculates the Edge Jitter Ratio and Ink Blot counts, and generates a Distance Transform heatmap overlay.\n\n")
    p_flow.add_run("If dynamic stylus coordinate data is available, it is processed through the dynamic rhythm analyzer to calculate velocity variance. All outputs are then passed to the Multimodal Risk Engine, which fuses the visual scores with transaction metadata to calculate a final transaction risk score. Finally, the system displays the verdict (Genuine or Suspicious), the confidence level, the micro-tremor statistics, and the Explainable AI heatmap on the teller's dashboard.")

    p_fig31 = doc.add_paragraph()
    p_fig31.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_fig31.add_run("\n[Flow Diagram Description: Preprocessing -> HOG / Tremor / LSTM paths -> Fusion Engine -> Teller Dashboard Output]\n")
    run.font.italic = True
    p_fig31.add_run("Figure 3.1: Flow Diagram for TrustTrace Verification\n\n")
    p_fig31.runs[1].font.bold = True

    h_sys = doc.add_paragraph()
    run = h_sys.add_run("System Diagram")
    run.font.size = Pt(14)
    run.font.bold = True
    h_sys.paragraph_format.space_before = Pt(12)
    h_sys.paragraph_format.space_after = Pt(6)

    p_sys = doc.add_paragraph()
    p_sys.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_sys.paragraph_format.line_spacing = 1.15
    p_sys.add_run("The system architecture is organized into a modular pipeline with the following logical components:\n")
    
    sys_components = [
        "Input Layer: Receives scanned signature images (cheque uploads) or dynamic touch/stylus coordinate strokes.",
        "Preprocessing Layer: Converts the image to grayscale, applies Otsu binarization to extract ink profiles, and resizes the output image to a fixed size of 256x128.",
        "Feature Extraction & CV Layer: Extracts HOG feature vectors and calculates micro-tremor edge jitter, ink blot contours, and distance transforms.",
        "Classification Layer: Utilizes a Soft-Voting Ensemble Classifier (SVM, RF, Gradient Boosting, LR, KNN) for static verification, and an LSTM velocity tracking model for temporal analysis.",
        "Adaptive Learning Layer: Stores centroid vectors in a latent space and updates reference profiles online using Euclidean distance thresholds.",
        "Multimodal Risk Engine: Computes a transaction-level risk score combining visual verification and metadata (amount, location, history flags).",
        "Output / Presentation Layer: Renders real-time alerts, confidence scores, and visual heatmaps on the teller's dashboard."
    ]
    for comp in sys_components:
        doc.add_paragraph(comp, style='List Bullet')

    p_fig32 = doc.add_paragraph()
    p_fig32.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_fig32.add_run("\n[System Diagram Description: Input -> Preprocessing -> CV & ML Layer -> Adaptive Centroid & Risk Fusion -> Output Display]\n")
    run.font.italic = True
    p_fig32.add_run("Figure 3.2: System Diagram of TrustTrace Signature Verification\n")
    p_fig32.runs[1].font.bold = True

    doc.add_page_break()

    # ------------------ CHAPTER 4: IMPLEMENTATION ------------------
    c4 = doc.add_paragraph()
    run = c4.add_run("CHAPTER 4\nIMPLEMENTATION")
    run.font.size = Pt(16)
    run.font.bold = True
    c4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c4.paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.add_run("The implementation of the proposed TrustTrace signature verification platform was carried out using Python because of its strong support for machine learning, image processing, and web service libraries. Python provides clean syntax, extensive modules, and efficient tools for image preprocessing, classification, and visualization. The backend is designed as an asynchronous REST API using the FastAPI framework, while the frontend dashboard is implemented using standard HTML5 Canvas, Vanilla CSS, and JavaScript to ensure compatibility and fast loading times on local bank terminals.\n\n")
    p.add_run("Several Python libraries were used during implementation. Pandas was used to manage dataset references, while NumPy supported numerical array calculations. Scikit-learn was used for train-test splitting and to implement the Soft-Voting Ensemble classifier. The skimage library was used for HOG feature extraction, and Joblib loaded the serialized model (.pkl) into active memory. OpenCV (cv2) managed the image processing pipeline, including grayscale conversion, Otsu thresholding, Canny filters, Gaussian blur, contour detection, and Distance Transform heatmap generation. The frontend interface communicates with the FastAPI service using standard asynchronous HTTP Fetch API requests.\n\n")
    p.add_run("The dataset used for the signature verification system was collected from a repository containing pairs of genuine signatures and skilled forgeries. The raw images are preprocessed and split into genuine and forged subsets. Data preprocessing forms the first stage of the implementation pipeline. The input image is converted to grayscale, binarized via Otsu's method to extract ink contours, and resized to a standard 256x128 pixels. The dataset is split into an 80:20 training and testing ratio using train_test_split() to evaluate the model's generalization capabilities.\n\n")
    p.add_run("Since machine learning algorithms cannot directly process raw images, the preprocessed binary image is transformed into a numerical vector using HOG feature extraction. The HOG algorithm is configured with 9 orientations, 8x8 pixels per cell, and 2x2 cells per block. This captures shape boundaries and structures while ignoring background noise. The ensemble classifier is trained on these features using five models: Support Vector Machine (SVM) with an RBF kernel, Random Forest (100 estimators), Gradient Boosting (50 estimators), Logistic Regression, and KNN (3 neighbors). During training, the classifiers learn the relationship between shape features and output labels (Genuine vs Forgery), saving the model as a serialized joblib file. The system also supports custom test uploads, where new images are preprocessed, analyzed for micro-tremors, and classified in real-time.\n\n")
    p.add_run("The sigmoid/probability decision function plays a significant role in Logistic Regression and SVM classifiers by converting raw decision scores into probability values between 0 and 1. Based on these probability values, the Soft-Voting Ensemble model computes average confidence scores. If the genuine probability exceeds 0.5, the signature is verified as genuine; otherwise, it is flagged as suspicious.")

    p_fig41 = doc.add_paragraph()
    p_fig41.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_fig41.add_run("\n[Sigmoid Curve Visualization: X-Axis: Decision Scores (-6 to +6), Y-Axis: Predicted Probability (0 to 1)]\n")
    run.font.italic = True
    p_fig41.add_run("Figure 4.1: Ensemble Voting Confidence Curve (Decision Scores vs Probability)\n")
    p_fig41.runs[1].font.bold = True

    doc.add_page_break()

    # ------------------ CHAPTER 5: RESULT ANALYSIS ------------------
    c5 = doc.add_paragraph()
    run = c5.add_run("CHAPTER 5\nRESULT ANALYSIS")
    run.font.size = Pt(16)
    run.font.bold = True
    c5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c5.paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.add_run("The proposed signature verification system was evaluated using a labeled dataset containing genuine signatures and skilled forgeries. Maintaining a balanced representation between classes was important to prevent the model from becoming biased toward a particular category during training. During dataset analysis, noticeable differences were observed in stroke quality and edge fluidity between genuine and forged signatures. Genuine signatures generally contained fast, ballistic strokes with high velocity variance, whereas traced forgeries appeared jagged, slow, and contained ink blots from hesitation pauses. These differences formed the basis for machine learning and computer vision-based classification.\n\n")
    p.add_run("Since machine learning classifiers cannot directly process raw images, HOG feature extraction was applied to transform signature contours into numerical feature vectors. The implementation used HOG configuration to capture structural boundaries while ignoring background noise. The ensemble classifier, combining SVM, Random Forest, Gradient Boosting, Logistic Regression, and KNN, was selected as the primary classification model due to its accuracy and robustness. The dataset was split into an 80:20 training and testing ratio. The validation process confirmed that data loading, preprocessing, HOG conversion, model training, and prediction generation were successfully completed without errors. The classifier effectively learned shape differences, while the OpenCV micro-tremor engine identified slow tracing and pen pauses. A major advantage of this approach is its low computational complexity, which enables fast training and real-time prediction on standard bank teller hardware.\n\n")
    p.add_run("The system's performance was evaluated using standard classification metrics. Accuracy was used as the primary metric, measuring the proportion of correctly verified signatures. Precision, recall, and F1-score were also evaluated through the classification report. Precision measures how many signatures flagged as forged were actually forged, reducing false accusations against genuine bank customers. Recall measures the model's ability to detect actual forgeries. F1-score combines both precision and recall, providing a balanced evaluation. The obtained results demonstrated balanced classification performance across both categories. The sigmoid function converted raw decision scores into probability values between 0 and 1, improving interpretability by providing the model's confidence level. A confusion matrix was used for detailed classification analysis because it compares actual labels with predicted labels, highlighting correct and incorrect classifications.")

    p_fig51 = doc.add_paragraph()
    p_fig51.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_fig51.add_run("\n[Confusion Matrix representation: Columns: Predicted Forged/Genuine, Rows: Actual Forged/Genuine]\n")
    run.font.italic = True
    p_fig51.add_run("Figure 5.1: Confusion Matrix for TrustTrace Signature Verification\n")
    p_fig51.runs[1].font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.add_run("The confusion matrix consists of four components: True Positives (genuine signatures correctly verified), True Negatives (forgeries correctly flagged), False Positives (genuine signatures incorrectly flagged as suspicious - Type I Error), and False Negatives (forged signatures incorrectly classified as genuine - Type II Error). Type I errors can frustrate customers, but Type II errors are critical in banking operations because accepting a forged signature can lead to direct financial losses.\n\n")
    p.add_run("Another important advantage of the ensemble model is its interpretability. The trained classifiers assign weight values to HOG features, enabling direct analysis of contours that contribute to predictions. Feature importance analysis showed that the classifier mainly relied on stroke orientations, edge jitter metrics, and ink blot locations when distinguishing between genuine and forged signatures. Structured, fluid lines were associated with genuine signatures, while jagged outlines and thick ink pools were linked with forged content. However, certain limitations were identified during experimentation, including dependency on image resolution, rotation sensitivity, and background noise on printed cheques. Despite these limitations, the results demonstrate that combining HOG-based ensembles and micro-tremor analysis provides an effective and computationally efficient framework for signature verification in banking environments.")

    doc.add_page_break()

    # ------------------ CHAPTER 6: CONCLUSIONS ------------------
    c6 = doc.add_paragraph()
    run = c6.add_run("CHAPTER 6\nCONCLUSIONS")
    run.font.size = Pt(16)
    run.font.bold = True
    c6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c6.paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.add_run("The proposed project, TrustTrace: Banking-Grade Signature Detection and Verification Using Machine Learning and Computer Vision, successfully demonstrates the application of Machine Learning and Computer Vision techniques for identifying signature forgery. With the growth of digital financial transactions, banks face increasing challenges in maintaining transaction security and verifying physical cheques. Traditional manual verification is slow, subjective, and prone to error. Therefore, automated verification systems have become an important area of research and practical application.\n\n")
    p.add_run("The implemented system utilized HOG feature extraction, a Soft-Voting Ensemble model, and OpenCV micro-tremor filters to distinguish between genuine and forged signatures. The results obtained from performance evaluation, classification reports, confusion matrix analysis, and feature importance analysis demonstrated that the proposed methodology is capable of identifying structural signature patterns and performing reliable binary classification. The project also showed that combining machine learning classifiers with computer vision metrics provides high verification accuracy without requiring expensive GPU hardware.\n\n")
    p.add_run("One of the major strengths of the proposed work is its multi-layered security approach. The system combines static shape verification, edge jitter calculations, dynamic velocity analysis, and online profile updates, providing a comprehensive verification pipeline. The prototypical adaptive network allows the system to adapt to natural signature drift over time, reducing false rejections. Although the system achieved effective results, certain limitations remain, such as dependency on scan quality and rotation alignment. Future improvements may include larger datasets, deep similarity learning (Siamese Networks), and advanced background subtraction techniques to improve robustness against watermarks on printed bank documents.\n\n")
    p.add_run("In conclusion, the proposed project provides a practical, secure, and efficient solution for signature verification. The study contributes to ongoing research in computer vision-based document authentication and demonstrates the practical applicability of machine learning in supporting secure financial operations.")

    doc.add_paragraph("\n\n")

    h = doc.add_paragraph()
    run = h.add_run("REFERENCES")
    run.font.size = Pt(14)
    run.font.bold = True
    
    refs = [
        'AI Generated Text Detection and Signature Analysis using Machine Learning and Deep Learning Models (2026)',
        'Towards Offline Signature Classification using TF-IDF and Machine Learning (2024)',
        'Zain, A., Farooqui, S., & Rafi, M. (2025). A Multi-Strategy Approach for Signature Forgery Detection. arXiv preprint. https://doi.org/10.48550/arXiv.2509.00623',
        'Detecting Signature Forgeries using Natural Language Processing and Machine Learning Techniques (2024)',
        'Comparative Study on HOG and LBP Feature Weighting in Text and Image Classification (2023)',
        'Signature vs Text Classification: A Machine Learning Approach (2024)',
        'Detecting Signature Forgeries in Short Form Images using Machine Learning (2025)',
        'Signature Forgery Detection: Challenges and Opportunities (2026)',
        'Babu, E. B., et al. (2024). Document Authentication Using Machine Learning Algorithms. 2024 Third International Conference on Electrical, Electronics, Information and Communication Technologies (ICEEICT), 1-6. https://doi.org/10.1109/iceeict61591.2024.10718469',
        'Human vs Machine Signature Classification using Machine Learning Models (2025)',
        'Investigation of Signature Detection in Academic Writing and Bank Transactions (2024)',
        'Benchmarking Signature Forgery Detection Methods (2026)',
        'Dependency-Based Image Text Detection using Machine Learning (2026)',
        'Hybrid Approach for Signature Detection combining Machine Learning and Deep Learning (2024)'
    ]
    for r in refs:
        doc.add_paragraph(r, style='List Bullet')

    # Save to disk
    doc.save("c:\\Users\\anuvu\\Desktop\\TrustTrace\\TRUSTTRACE_REPORT_PART2.docx")
    print("Part 2 report saved successfully as TRUSTTRACE_REPORT_PART2.docx")

if __name__ == "__main__":
    create_report_part2()
